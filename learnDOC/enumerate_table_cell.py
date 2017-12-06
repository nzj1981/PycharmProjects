#!/usr/bin/env python3.6
# coding=utf-8
import docx
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE

# enumerate a cell of table
document = docx.Document('demo2.doc')

table = document.tables[0]

for row, obj_row in enumerate(table.rows):
    for col, cell in enumerate(obj_row.cells):
        cell.text = '{}, {}'.format(row, col)
# add width of table column 4
table.columns[4].width = Cm(4)
# output table all style
styles = document.styles
table_style = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE]
for style in table_style:
    print('table all style:{}'.format(style.name))
# update table style
table.style = 'Medium Grid 2 Accent 5'
# update text of cell at 1 rows
row = table.rows[1]
row.cells[0].text = '流金岁月'
row.cells[1].text = '水调歌头'
# Statistics row and column total
row_cn = len(table.rows)
col_cn = len(table.columns)
print('row = {}, column = {}'.format(row_cn, col_cn))
# add a line to the table
# row = table.add_row()
# Adding a picture
document.add_page_break()
document.add_picture('pic.jpg', width=Cm(3))
document.save('demo2.doc')
