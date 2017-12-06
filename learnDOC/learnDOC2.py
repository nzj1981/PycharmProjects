#!/usr/bin/env python3.6
# coding=utf-8

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

# open new document
document = docx.Document()
# add a paragraph
paragraph = document.add_paragraph('作者： 苏轼', 'Subtitle')
prior_paragraph = paragraph.insert_paragraph_before(u'水调歌头', 'Title')
document.add_paragraph(
 '''
明月几时有，把酒问青天。
不知天上宫阙，今夕是何年？
我欲乘风归去，又恐琼楼玉宇，
高处不胜寒。
起舞弄清影，何似在人间！

转朱阁，低绮户，照无眠。
不应有恨，何事长向别时圆？
人有悲欢离合，月有阴晴圆缺，
此事古难全。
但愿人长久，千里共婵娟。
''', 'Body Text 2')
# Adding a heading
document.add_heading('The Real meaning of the universe')
document.add_heading('标题0', level=0)
document.add_heading('标题1', level=1)
document.add_heading('标题2', level=2)
document.add_heading('标题3', level=3)
# View paragraph number of words and the use of style
for p in document.paragraphs:
    print('该段有{}字'.format(len(p.text)))
    print('该段所使用样式为{}'.format(p.style.name))
# The Styles object is also iterable. By using the identification properties on BaseStyle, various subsets of the
# defined styles can be generated.
styles = document.styles
print('\n'.join([s.name for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]))
paragraph2 = document.add_paragraph('流金岁月')
paragraph2.style = 'List Bullet 2'
paragraph3 = document.add_paragraph('水调歌头')
paragraph3.style = 'List Number 2'
# character style
paragraph4 = document.add_paragraph()
paragraph4.add_run(u'《水调歌头·游泳》,')
paragraph4.add_run(u'才饮长沙水， ').bold = True
paragraph4.add_run(u'又食武昌鱼。').italic = True
run = paragraph4.add_run(u'万里长江横渡，')
run.style = 'Intense Reference'
# Set the font and font size
run.font.name = 'COOPBL'
run.font.size = Pt(20)
# Adding a page bread
document.add_page_break()
# Adding a tables
document.add_paragraph()
table = document.add_table(rows=9, cols = 10, style='Colorful Shading')
cell_1 = table.cell(1, 2)
cell_2 = table.cell(4, 7)
cell_1.merge(cell_2)
# save document
document.save('demo2.doc')
