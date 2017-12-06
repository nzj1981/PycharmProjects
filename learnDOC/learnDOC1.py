#!/usr/bin/env python3.6
# coding=utf-8

import docx
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

# open new document
document = docx.Document()

# Add different levels of title
document.add_heading('Document Title', 0)
document.add_heading(u'二级标题', 1)
document.add_heading(u'  三级标题', 2)
# add text
paragraph = document.add_paragraph(u'\n添加文本')
# set size of font
run = paragraph.add_run(u'\n设置字号')
run.font.size = Pt(24)
# set font
run = paragraph.add_run(u'\n设置字体')
run.font.name = 'Consolas'
# set chinese font
run = paragraph.add_run(u'\n设置中文字体')
run.font.name = u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
# set italic
run = paragraph.add_run(u'\n设置斜体')
run.italic = True
run.font.name = u'宋体'
r1 = run._element
r1.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
# set bold
run = paragraph.add_run(u'\n设置粗体')
run.bold = True
run.font.name = u'方正楷体'
r2 = run._element
r2.rPr.rFonts.set(qn('w:eastAsia'), u'方正楷体')
# add quote
document.add_paragraph(u'\n增加引用', style='Intense Quote')
# Add an ordered list
document.add_paragraph(
    u'有序列表元素1', style='List Number'
)
document.add_paragraph(
    u'有序列表元素2', style='List Number'
)
# Add unordered list
document.add_paragraph(
    u'无序列表元素1', style='List Bullet'
)
document.add_paragraph(
    u'无序列表元素2', style='List Bullet'
)
# add picture
document.add_picture('pic.jpg', width=Inches(1.25))
# add table
table = document.add_table(rows=3, cols=3, style='Colorful Shading')
t_cells = table.rows[0].cells
t_cells[0].text = '第一行第一列'
t_cells[1].text = '第一行第二列'
t_cells[2].text = '第一行第三列'

t_cells = table.rows[1].cells
t_cells[0].text = '第二行第一列'
t_cells[1].text = '第二行第二列'
t_cells[2].text = '第二行第三列'
# Add page
document.add_page_break()

# save file
document.save('demo.doc')
